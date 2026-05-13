"""Shared pytest fixtures and helpers.

A deterministic ``HashEmbedder`` lives here so any test that needs vector
embeddings can use it without spinning up a real sentence-transformer model
(saves ~400MB of downloads + ~1s startup). The hash-based scheme is *not*
semantically meaningful — paraphrase neighbours won't surface — but identical
strings always produce identical vectors, which is enough for round-trip and
schema tests.

Session-scoped ``sample_pdf_path`` and ``sample_docx_path`` fixtures generate
small valid PDF and DOCX files under ``tmp_path_factory`` so binary fixtures
don't have to be checked in. ``reportlab`` and ``python-docx`` are dev deps.
"""

from __future__ import annotations

from hashlib import sha256
from pathlib import Path
from typing import TYPE_CHECKING

import numpy as np
import pytest

if TYPE_CHECKING:
    from numpy.typing import NDArray


class HashEmbedder:
    """Hash-derived deterministic embeddings for hermetic tests."""

    def __init__(self, dim: int = 64) -> None:
        if dim % 4 != 0:
            raise ValueError("dim must be a multiple of 4")
        self._dim = dim

    @property
    def dim(self) -> int:
        return self._dim

    def embed_texts(self, texts: list[str]) -> NDArray[np.float32]:
        out = np.zeros((len(texts), self._dim), dtype=np.float32)
        for i, text in enumerate(texts):
            data = sha256(text.encode("utf-8")).digest()
            while len(data) < self._dim * 4:
                data = data + sha256(data).digest()
            # uint32 → float64 in [-1, 1] avoids NaN/inf bit patterns that
            # raw float32 reinterpretation would produce.
            raw = np.frombuffer(data[: self._dim * 4], dtype=np.uint32).astype(np.float64)
            arr = (raw / np.iinfo(np.uint32).max) * 2.0 - 1.0
            norm = float(np.linalg.norm(arr))
            if norm > 0.0:
                arr = arr / norm
            out[i] = arr.astype(np.float32)
        return out


# --- generated binary fixtures ----------------------------------------------


@pytest.fixture(scope="session")
def sample_pdf_path(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """Generate a small two-line PDF under tmp; cached for the session."""
    from reportlab.pdfgen import canvas

    out = tmp_path_factory.mktemp("binary_fixtures") / "sample.pdf"
    pdf = canvas.Canvas(str(out))
    pdf.drawString(100, 750, "Sample PDF heading line.")
    pdf.drawString(100, 730, "First body sentence about widgets.")
    pdf.drawString(100, 710, "Second body sentence about gadgets.")
    pdf.showPage()
    pdf.save()
    return out


@pytest.fixture(scope="session")
def sample_docx_path(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """Generate a DOCX with heading + paragraphs + 2x2 table; cached for the session."""
    from docx import Document as DocxDocument

    out = tmp_path_factory.mktemp("binary_fixtures") / "sample.docx"
    doc = DocxDocument()
    doc.add_heading("Sample DOCX title", level=1)
    doc.add_paragraph("First DOCX body paragraph.")
    doc.add_paragraph("Second DOCX body paragraph.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A1"
    table.rows[0].cells[1].text = "B1"
    table.rows[1].cells[0].text = "A2"
    table.rows[1].cells[1].text = "B2"
    doc.save(str(out))
    return out
